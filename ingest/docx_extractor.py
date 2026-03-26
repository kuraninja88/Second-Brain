"""
Extract text chunks from DOCX files.

Modes (determined by source_type):
  paper_note   — per-paper reading notes (Bacaan/Catatan/)
  daily_note   — daily research notes (Catatan/)
  thesis_draft — thesis draft sections (Thesis/)
  lit_notes    — structured lit review notes (lit_review_notes_ESG_MA_SCI_v2.docx)
  general      — any other DOCX
"""
import re
from pathlib import Path
from typing import Optional
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from config import CHUNK_SIZE, CHUNK_OVERLAP, MIN_CHUNK_LEN


def _safe_text(para) -> str:
    try:
        return para.text.encode("utf-8", errors="replace").decode("utf-8").strip()
    except Exception:
        return ""


def _chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks by approximate character count."""
    # rough chars-per-token ≈ 4
    size = CHUNK_SIZE * 4
    overlap = CHUNK_OVERLAP * 4

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        chunk = text[start:end].strip()
        if len(chunk) >= MIN_CHUNK_LEN:
            chunks.append(chunk)
        start += size - overlap
    return chunks


def _parse_date_from_filename(stem: str) -> str:
    """Try to extract a date string from 'N. Catatan DD MMM YYYY'."""
    m = re.search(r"Catatan\s+(.+)", stem, re.IGNORECASE)
    return m.group(1).strip() if m else ""


def extract_docx(file_path: Path, source_type: str) -> list[dict]:
    """
    Returns list of chunk dicts:
        {"text": ..., "metadata": {...}}
    """
    from docx import Document
    try:
        doc = Document(str(file_path))
    except Exception as e:
        print(f"  [WARN] Cannot open {file_path.name}: {e}")
        return []

    stem = file_path.stem
    date_str = _parse_date_from_filename(stem) if "catatan" in stem.lower() else ""

    base_meta = {
        "source_type": source_type,
        "file_path":   str(file_path),
        "filename":    file_path.name,
        "date":        date_str,
    }

    # ── Structured lit review notes: split on numbered headings ──────────────
    if source_type == "lit_notes":
        return _extract_lit_notes(doc, file_path, base_meta)

    # ── Thesis drafts: split on Word heading styles ───────────────────────────
    if source_type == "thesis_draft":
        return _extract_by_headings(doc, file_path, base_meta)

    # ── Default: chunk all paragraphs together ────────────────────────────────
    full_text = "\n".join(_safe_text(p) for p in doc.paragraphs if _safe_text(p))
    chunks = _chunk_text(full_text)
    return [
        {"text": c, "metadata": {**base_meta, "chunk_index": i}}
        for i, c in enumerate(chunks)
    ]


def _extract_lit_notes(doc, file_path: Path, base_meta: dict) -> list[dict]:
    """Split on patterns like '1. Author (YYYY)' or '## Author'."""
    heading_re = re.compile(r"^\d+\.\s+\w|^#{1,3}\s+\w", re.MULTILINE)
    results = []
    current_heading = "Introduction"
    current_paras: list[str] = []

    for para in doc.paragraphs:
        text = _safe_text(para)
        if not text:
            continue
        if heading_re.match(text) or (para.style is not None and para.style.name.startswith("Heading") and len(text) < 120):
            # Save previous section
            if current_paras:
                section_text = "\n".join(current_paras)
                for i, chunk in enumerate(_chunk_text(section_text)):
                    results.append({
                        "text": chunk,
                        "metadata": {
                            **base_meta,
                            "section": current_heading,
                            "chunk_index": i,
                        }
                    })
            current_heading = text
            current_paras = []
        else:
            current_paras.append(text)

    # Last section
    if current_paras:
        section_text = "\n".join(current_paras)
        for i, chunk in enumerate(_chunk_text(section_text)):
            results.append({
                "text": chunk,
                "metadata": {**base_meta, "section": current_heading, "chunk_index": i}
            })
    return results


def _extract_by_headings(doc, file_path: Path, base_meta: dict) -> list[dict]:
    """Split thesis drafts by Word heading styles."""
    results = []
    current_heading = file_path.stem
    current_paras: list[str] = []

    for para in doc.paragraphs:
        text = _safe_text(para)
        if not text:
            continue
        if para.style is not None and para.style.name.startswith("Heading"):
            if current_paras:
                section_text = "\n".join(current_paras)
                for i, chunk in enumerate(_chunk_text(section_text)):
                    results.append({
                        "text": chunk,
                        "metadata": {
                            **base_meta,
                            "section": current_heading,
                            "chunk_index": i,
                        }
                    })
            current_heading = text
            current_paras = []
        else:
            current_paras.append(text)

    if current_paras:
        section_text = "\n".join(current_paras)
        for i, chunk in enumerate(_chunk_text(section_text)):
            results.append({
                "text": chunk,
                "metadata": {**base_meta, "section": current_heading, "chunk_index": i}
            })

    # Fallback: no headings found → treat as one block
    if not results:
        full = "\n".join(_safe_text(p) for p in doc.paragraphs if _safe_text(p))
        for i, chunk in enumerate(_chunk_text(full)):
            results.append({"text": chunk, "metadata": {**base_meta, "chunk_index": i}})

    return results
